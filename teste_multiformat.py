"""
Script de teste rapido para o multi-formato (Caxinguele v2)
Cria documentos de teste em varios formatos e processa cada um.

Uso:
    python teste_multiformat.py
    python teste_multiformat.py --formato epub
    python teste_multiformat.py --formato txt --tts  (converte para audio tambem)
"""

import sys
import argparse
import logging
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")


def criar_txt_teste(pasta: Path) -> Path:
    """Cria arquivo TXT de teste"""
    caminho = pasta / "teste.txt"
    caminho.write_text(
        "Este eh um teste do sistema Caxinguele v2.\n\n"
        "O sistema agora suporta multiplos formatos de documento.\n\n"
        "Voce pode converter PDFs, Word, EPUB, emails e muito mais.\n\n"
        "Todos os documentos ficam disponiveis na Alexa para leitura por voz.\n\n"
        "O sistema e acessivel para usuarios com deficiencia visual.",
        encoding="utf-8"
    )
    return caminho


def criar_md_teste(pasta: Path) -> Path:
    """Cria arquivo Markdown de teste"""
    caminho = pasta / "teste.md"
    caminho.write_text(
        "# Teste do Caxinguele v2\n\n"
        "## Sobre o Sistema\n\n"
        "Este e um **teste** do sistema de audiobooks.\n\n"
        "## Funcionalidades\n\n"
        "- Suporte multi-formato\n"
        "- Drag-and-drop\n"
        "- Classificacao automatica\n\n"
        "## Conclusao\n\n"
        "O sistema funciona muito bem para acessibilidade.",
        encoding="utf-8"
    )
    return caminho


def criar_html_teste(pasta: Path) -> Path:
    """Cria arquivo HTML de teste"""
    caminho = pasta / "teste.html"
    caminho.write_text(
        """<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <title>Noticia de Teste</title>
</head>
<body>
    <h1>Sistema de Audiobooks Lanca Suporte Multi-Formato</h1>
    <p>O Projeto Caxinguele lanca hoje a versao 2.0 com suporte a
    mais de 19 tipos de documento, incluindo Word, EPUB, emails e imagens.</p>
    <p>A nova versao inclui deteccao automatica de tipo, drag-and-drop
    e organizacao automatica no Google Drive por categoria.</p>
    <p>O sistema foi desenvolvido especialmente para acessibilidade,
    permitindo que usuarios com deficiencia visual acessem qualquer
    documento por voz via Alexa.</p>
</body>
</html>""",
        encoding="utf-8"
    )
    return caminho


def criar_eml_teste(pasta: Path) -> Path:
    """Cria arquivo EML de teste"""
    caminho = pasta / "teste.eml"
    caminho.write_text(
        "From: andre@exemplo.com\n"
        "To: amigo@exemplo.com\n"
        "Subject: Teste do Sistema Caxinguele\n"
        "Date: Thu, 19 Feb 2026 10:00:00 +0000\n"
        "Content-Type: text/plain; charset=utf-8\n"
        "\n"
        "Ola!\n\n"
        "Este e um email de teste para o sistema Caxinguele v2.\n\n"
        "O sistema pode converter emails em audiobooks para a Alexa.\n\n"
        "Abracos,\nAndre",
        encoding="utf-8"
    )
    return caminho


def criar_docx_teste(pasta: Path) -> Path:
    """Cria arquivo DOCX de teste"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.core_properties.title = "Documento de Teste"
        doc.core_properties.author = "Sistema Caxinguele"

        # Capitulo 1
        doc.add_heading("Introducao ao Caxinguele v2", level=1)
        doc.add_paragraph(
            "O Projeto Caxinguele e um sistema de conversao de documentos "
            "em audiobooks acessiveis via Amazon Alexa."
        )
        doc.add_paragraph(
            "O sistema suporta mais de 19 formatos de documento, incluindo "
            "PDF, Word, EPUB, emails e imagens com OCR."
        )

        # Capitulo 2
        doc.add_heading("Funcionalidades Principais", level=1)
        doc.add_paragraph(
            "A versao 2 traz suporte multi-formato completo, deteccao automatica "
            "de tipo de documento, drag-and-drop na interface e organizacao "
            "automatica por categoria no Google Drive."
        )

        # Capitulo 3
        doc.add_heading("Acessibilidade", level=1)
        doc.add_paragraph(
            "O sistema foi desenvolvido com foco em acessibilidade. "
            "Usuarios com deficiencia visual podem acessar qualquer documento "
            "por comandos de voz na Alexa, sem necessidade de ver a tela."
        )

        caminho = pasta / "teste.docx"
        doc.save(str(caminho))
        return caminho

    except ImportError:
        logging.warning("python-docx nao disponivel, pulando teste DOCX")
        return None


def testar_formato(caminho: Path, fazer_tts: bool = False) -> bool:
    """Testa processamento de um arquivo"""
    from doc_processor import processar_documento
    from doc_classifier import classificar_documento

    if not caminho or not caminho.exists():
        return False

    print(f"\n{'='*50}")
    print(f"Testando: {caminho.name}")
    print(f"{'='*50}")

    try:
        # Processa documento
        livro = processar_documento(caminho)
        total_palavras = sum(c.palavras for c in livro.capitulos)

        print(f"  Titulo    : {livro.titulo}")
        print(f"  Capitulos : {len(livro.capitulos)}")
        print(f"  Palavras  : {total_palavras}")
        if livro.autor:
            print(f"  Autor     : {livro.autor}")

        # Classifica tipo
        texto_amostra = livro.capitulos[0].texto[:1000] if livro.capitulos else ""
        classif = classificar_documento(caminho, texto_amostra, livro.num_paginas)
        print(f"  Tipo      : {classif.icone} {classif.nome} ({classif.confianca:.0%}, via {classif.metodo})")
        print(f"  Pasta Drive: {classif.pasta_drive}")

        # Preview do texto
        if livro.capitulos:
            preview = livro.capitulos[0].texto[:200].replace('\n', ' ')
            print(f"  Preview   : {preview}...")

        # Converte para TTS se solicitado
        if fazer_tts:
            print(f"\n  Iniciando conversao TTS...")
            from tts_engine import converter_texto_para_audio
            from config import VOZES_PT_BR
            from pathlib import Path as P

            pasta_saida = P("temp_teste")
            pasta_saida.mkdir(exist_ok=True)

            for cap in livro.capitulos[:2]:  # Apenas primeiros 2 capitulos
                if cap.palavras > 0:
                    arquivo_mp3 = pasta_saida / f"{caminho.stem}_cap{cap.numero:02d}.mp3"
                    sucesso = converter_texto_para_audio(
                        cap.texto[:500],  # Apenas 500 chars para teste rapido
                        arquivo_mp3,
                        voz=VOZES_PT_BR["thalita"]
                    )
                    if sucesso:
                        tamanho = arquivo_mp3.stat().st_size / 1024
                        print(f"  TTS Cap {cap.numero}: OK ({tamanho:.0f} KB) -> {arquivo_mp3}")
                    else:
                        print(f"  TTS Cap {cap.numero}: FALHOU")

        print(f"  RESULTADO : OK")
        return True

    except Exception as e:
        print(f"  RESULTADO : ERRO - {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description="Teste multi-formato Caxinguele v2")
    parser.add_argument("--formato", choices=["txt", "md", "html", "eml", "docx", "todos"],
                        default="todos", help="Formato a testar")
    parser.add_argument("--tts", action="store_true",
                        help="Converte para audio alem de extrair texto")
    args = parser.parse_args()

    logging.basicConfig(level=logging.WARNING)  # Silencia logs desnecessarios

    # Pasta temporaria para arquivos de teste
    pasta_teste = Path("temp_testes")
    pasta_teste.mkdir(exist_ok=True)

    print("=" * 50)
    print("TESTE MULTI-FORMATO - Caxinguele v2")
    print("=" * 50)

    # Cria e testa cada formato
    formatos = {
        "txt":  (criar_txt_teste, "TXT"),
        "md":   (criar_md_teste, "Markdown"),
        "html": (criar_html_teste, "HTML"),
        "eml":  (criar_eml_teste, "Email EML"),
        "docx": (criar_docx_teste, "Word DOCX"),
    }

    if args.formato == "todos":
        selecionados = list(formatos.keys())
    else:
        selecionados = [args.formato]

    resultados = {}
    for fmt in selecionados:
        criador, nome = formatos[fmt]
        arquivo = criador(pasta_teste)
        sucesso = testar_formato(arquivo, fazer_tts=args.tts)
        resultados[nome] = sucesso

    # Resumo
    print(f"\n{'='*50}")
    print("RESUMO DOS TESTES")
    print(f"{'='*50}")
    for nome, sucesso in resultados.items():
        status = "OK" if sucesso else "FALHOU"
        print(f"  {nome:<20}: {status}")

    total_ok = sum(1 for s in resultados.values() if s)
    print(f"\nTotal: {total_ok}/{len(resultados)} formatos OK")

    # Limpa pasta de teste
    import shutil
    shutil.rmtree(pasta_teste, ignore_errors=True)


if __name__ == "__main__":
    main()
