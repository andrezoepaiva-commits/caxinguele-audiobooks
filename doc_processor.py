"""
Módulo para processamento de documentos multi-formato
Suporta: PDF, DOCX, RTF, ODT, TXT, MD, EPUB, MOBI, EML, MSG, HTML, Imagens (OCR)

Reutiliza as classes LivroProcessado e Capitulo do pdf_processor.py
"""

import logging
import re
from pathlib import Path
from typing import List, Optional

from pdf_processor import LivroProcessado, Capitulo, processar_pdf, limpar_texto_pdf
from utils import limpar_texto_para_tts


# ==================== FORMATOS SUPORTADOS ====================

# Mapeamento extensão → função de extração
FORMATOS_SUPORTADOS = {
    # Documentos de texto
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.rtf': 'rtf',
    '.odt': 'odt',
    '.txt': 'txt',
    '.md': 'markdown',
    # Livros digitais
    '.epub': 'epub',
    '.mobi': 'mobi',
    # Email
    '.eml': 'eml',
    '.msg': 'msg',
    # Web
    '.html': 'html',
    '.htm': 'html',
    # Imagens (OCR)
    '.jpg': 'imagem',
    '.jpeg': 'imagem',
    '.png': 'imagem',
    '.tiff': 'imagem',
    '.tif': 'imagem',
    '.bmp': 'imagem',
    '.webp': 'imagem',
}

# Extensões para o filtro do seletor de arquivos
FILTRO_EXTENSOES = [
    ("Todos os documentos", "*.pdf *.docx *.rtf *.odt *.txt *.md *.epub *.mobi *.eml *.msg *.html *.htm *.jpg *.jpeg *.png *.tiff *.tif *.bmp *.webp"),
    ("PDF", "*.pdf"),
    ("Word", "*.docx"),
    ("eBook", "*.epub *.mobi"),
    ("Texto", "*.txt *.md *.rtf *.odt"),
    ("Email", "*.eml *.msg"),
    ("Web", "*.html *.htm"),
    ("Imagem (OCR)", "*.jpg *.jpeg *.png *.tiff *.tif *.bmp *.webp"),
    ("Todos", "*.*"),
]


# ==================== FUNÇÃO PRINCIPAL ====================

def processar_documento(caminho: Path) -> LivroProcessado:
    """
    Processa qualquer documento suportado e retorna LivroProcessado.
    Detecta o formato pela extensão e chama a função adequada.

    Args:
        caminho: Path do arquivo

    Returns:
        LivroProcessado com capítulos extraídos
    """
    caminho = Path(caminho)

    if not caminho.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {caminho}")

    extensao = caminho.suffix.lower()

    if extensao not in FORMATOS_SUPORTADOS:
        raise ValueError(
            f"Formato não suportado: {extensao}\n"
            f"Formatos aceitos: {', '.join(sorted(FORMATOS_SUPORTADOS.keys()))}"
        )

    tipo = FORMATOS_SUPORTADOS[extensao]
    logging.info(f"Processando {caminho.name} (formato: {tipo})")

    # PDF usa o processador original (já maduro e testado)
    if tipo == 'pdf':
        return processar_pdf(caminho)

    # Demais formatos
    extratores = {
        'docx': _extrair_docx,
        'rtf': _extrair_rtf,
        'odt': _extrair_odt,
        'txt': _extrair_txt,
        'markdown': _extrair_markdown,
        'epub': _extrair_epub,
        'mobi': _extrair_mobi,
        'eml': _extrair_eml,
        'msg': _extrair_msg,
        'html': _extrair_html,
        'imagem': _extrair_imagem_ocr,
    }

    extrator = extratores[tipo]
    return extrator(caminho)


def formato_suportado(caminho: Path) -> bool:
    """Verifica se o formato do arquivo é suportado"""
    return Path(caminho).suffix.lower() in FORMATOS_SUPORTADOS


# ==================== EXTRATORES POR FORMATO ====================

def _criar_livro_simples(caminho: Path, titulo: str, texto: str) -> LivroProcessado:
    """
    Helper: cria LivroProcessado com um único capítulo (para documentos curtos).
    Se o texto for longo, divide em capítulos artificiais.
    """
    texto = limpar_texto_para_tts(texto)
    palavras = len(texto.split())

    livro = LivroProcessado(caminho)
    livro.titulo = titulo
    livro.autor = ""
    livro.num_paginas = max(1, palavras // 300)  # Estimativa ~300 palavras/página
    livro.tamanho_mb = caminho.stat().st_size / (1024 * 1024)

    if palavras <= 5000:
        # Documento curto: um único capítulo
        capitulo = Capitulo(
            numero=1,
            titulo=titulo,
            texto=texto,
            pagina_inicio=0
        )
        livro.capitulos = [capitulo]
    else:
        # Documento longo: divide em partes de ~3000 palavras
        livro.capitulos = _dividir_texto_em_capitulos(texto, titulo)

    total_palavras = sum(c.palavras for c in livro.capitulos)
    logging.info(f"Extraído: {len(livro.capitulos)} capítulo(s), {total_palavras:,} palavras")
    return livro


def _dividir_texto_em_capitulos(texto: str, titulo_base: str, max_palavras: int = 3000) -> List[Capitulo]:
    """Divide texto longo em capítulos artificiais por parágrafos"""
    paragrafos = texto.split('\n\n')
    capitulos = []
    chunk_atual = []
    palavras_atual = 0
    num = 1

    for paragrafo in paragrafos:
        palavras_p = len(paragrafo.split())
        if palavras_atual + palavras_p > max_palavras and chunk_atual:
            capitulos.append(Capitulo(
                numero=num,
                titulo=f"{titulo_base} - Parte {num}",
                texto='\n\n'.join(chunk_atual),
                pagina_inicio=0
            ))
            num += 1
            chunk_atual = [paragrafo]
            palavras_atual = palavras_p
        else:
            chunk_atual.append(paragrafo)
            palavras_atual += palavras_p

    if chunk_atual:
        capitulos.append(Capitulo(
            numero=num,
            titulo=f"{titulo_base} - Parte {num}" if num > 1 else titulo_base,
            texto='\n\n'.join(chunk_atual),
            pagina_inicio=0
        ))

    return capitulos


# ---- DOCX (Word) ----

def _extrair_docx(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo Word (.docx)"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instale python-docx: pip install python-docx")

    doc = Document(str(caminho))
    titulo = caminho.stem.replace("_", " ").replace("-", " ")

    # Tenta extrair título do documento
    if doc.core_properties.title:
        titulo = doc.core_properties.title

    # Tenta detectar capítulos por headings
    capitulos = []
    texto_atual = []
    titulo_cap = None
    num_cap = 0

    for para in doc.paragraphs:
        # Heading = novo capítulo
        if para.style.name.startswith('Heading') and para.text.strip():
            if texto_atual and titulo_cap:
                num_cap += 1
                texto_junto = '\n\n'.join(texto_atual)
                texto_junto = limpar_texto_para_tts(texto_junto)
                if len(texto_junto.split()) > 10:  # Ignora capítulos vazios
                    capitulos.append(Capitulo(
                        numero=num_cap,
                        titulo=titulo_cap,
                        texto=texto_junto,
                        pagina_inicio=0
                    ))
                texto_atual = []
            titulo_cap = para.text.strip()
        elif para.text.strip():
            texto_atual.append(para.text.strip())

    # Último capítulo
    if texto_atual:
        num_cap += 1
        texto_junto = '\n\n'.join(texto_atual)
        texto_junto = limpar_texto_para_tts(texto_junto)
        if not titulo_cap:
            titulo_cap = titulo
        if len(texto_junto.split()) > 10:
            capitulos.append(Capitulo(
                numero=num_cap,
                titulo=titulo_cap,
                texto=texto_junto,
                pagina_inicio=0
            ))

    # Se não encontrou capítulos por headings, usa texto corrido
    if not capitulos:
        texto_completo = '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        return _criar_livro_simples(caminho, titulo, texto_completo)

    livro = LivroProcessado(caminho)
    livro.titulo = titulo
    livro.autor = doc.core_properties.author or ""
    livro.num_paginas = len(doc.paragraphs) // 30  # Estimativa
    livro.tamanho_mb = caminho.stat().st_size / (1024 * 1024)
    livro.capitulos = capitulos

    total_palavras = sum(c.palavras for c in capitulos)
    logging.info(f"DOCX extraído: {len(capitulos)} capítulo(s), {total_palavras:,} palavras")
    return livro


# ---- RTF ----

def _extrair_rtf(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo RTF"""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise ImportError("Instale striprtf: pip install striprtf")

    conteudo_rtf = caminho.read_text(encoding='utf-8', errors='replace')
    texto = rtf_to_text(conteudo_rtf)
    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    return _criar_livro_simples(caminho, titulo, texto)


# ---- ODT (LibreOffice) ----

def _extrair_odt(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo ODT (OpenDocument)"""
    try:
        from odf.opendocument import load
        from odf.text import P
        from odf import teletype
    except ImportError:
        raise ImportError("Instale odfpy: pip install odfpy")

    doc = load(str(caminho))
    paragrafos = doc.getElementsByType(P)
    texto = '\n\n'.join(teletype.extractText(p) for p in paragrafos if teletype.extractText(p).strip())
    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    return _criar_livro_simples(caminho, titulo, texto)


# ---- TXT ----

def _extrair_txt(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo TXT"""
    # Tenta diferentes encodings
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            texto = caminho.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        texto = caminho.read_text(encoding='utf-8', errors='replace')

    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    return _criar_livro_simples(caminho, titulo, texto)


# ---- Markdown ----

def _extrair_markdown(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo Markdown, removendo formatação"""
    texto = caminho.read_text(encoding='utf-8', errors='replace')
    titulo = caminho.stem.replace("_", " ").replace("-", " ")

    # Remove formatação markdown básica
    texto = re.sub(r'^#{1,6}\s+', '', texto, flags=re.MULTILINE)  # Headers
    texto = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', texto)  # Links
    texto = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', texto)  # Imagens
    texto = re.sub(r'```[\s\S]*?```', '', texto)  # Blocos de código
    texto = re.sub(r'`([^`]+)`', r'\1', texto)  # Código inline
    texto = re.sub(r'^[-*+]\s+', '', texto, flags=re.MULTILINE)  # Listas
    texto = re.sub(r'^\d+\.\s+', '', texto, flags=re.MULTILINE)  # Listas numeradas
    texto = re.sub(r'^>\s+', '', texto, flags=re.MULTILINE)  # Blockquotes
    texto = re.sub(r'\*\*([^*]+)\*\*', r'\1', texto)  # Bold
    texto = re.sub(r'\*([^*]+)\*', r'\1', texto)  # Italic

    return _criar_livro_simples(caminho, titulo, texto)


# ---- EPUB ----

def _extrair_epub(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo EPUB"""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("Instale ebooklib e beautifulsoup4: pip install ebooklib beautifulsoup4")

    book = epub.read_epub(str(caminho))

    # Extrai título
    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    if book.get_metadata('DC', 'title'):
        titulo = book.get_metadata('DC', 'title')[0][0]

    # Extrai autor
    autor = ""
    if book.get_metadata('DC', 'creator'):
        autor = book.get_metadata('DC', 'creator')[0][0]

    # Extrai capítulos do spine (ordem de leitura)
    capitulos = []
    num_cap = 0

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        html_content = item.get_content().decode('utf-8', errors='replace')
        soup = BeautifulSoup(html_content, 'html.parser')

        # Tenta pegar título do capítulo
        titulo_cap = None
        for tag in ['h1', 'h2', 'h3', 'title']:
            elem = soup.find(tag)
            if elem and elem.get_text(strip=True):
                titulo_cap = elem.get_text(strip=True)
                break

        # Extrai texto
        texto = soup.get_text(separator='\n\n', strip=True)
        texto = limpar_texto_para_tts(texto)

        if len(texto.split()) < 20:
            continue  # Pula capítulos muito curtos (capa, sumário, etc.)

        num_cap += 1
        if not titulo_cap:
            titulo_cap = f"Capítulo {num_cap}"

        capitulos.append(Capitulo(
            numero=num_cap,
            titulo=titulo_cap,
            texto=texto,
            pagina_inicio=0
        ))

    if not capitulos:
        raise ValueError(f"Não foi possível extrair texto do EPUB: {caminho.name}")

    livro = LivroProcessado(caminho)
    livro.titulo = titulo
    livro.autor = autor
    livro.num_paginas = sum(c.palavras for c in capitulos) // 300
    livro.tamanho_mb = caminho.stat().st_size / (1024 * 1024)
    livro.capitulos = capitulos

    total_palavras = sum(c.palavras for c in capitulos)
    logging.info(f"EPUB extraído: {len(capitulos)} capítulo(s), {total_palavras:,} palavras")
    return livro


# ---- MOBI ----

def _extrair_mobi(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo MOBI (Kindle)"""
    try:
        import mobi
    except ImportError:
        raise ImportError("Instale mobi: pip install mobi")

    # mobi extrai para HTML, depois processamos como HTML
    tempdir, filepath = mobi.extract(str(caminho))
    filepath = Path(filepath)

    if filepath.suffix.lower() in ('.html', '.htm'):
        livro = _extrair_html(filepath)
    else:
        # Se extraiu como texto
        texto = filepath.read_text(encoding='utf-8', errors='replace')
        titulo = caminho.stem.replace("_", " ").replace("-", " ")
        livro = _criar_livro_simples(caminho, titulo, texto)

    # Corrige caminho e título
    livro.caminho = caminho
    livro.titulo = caminho.stem.replace("_", " ").replace("-", " ")

    # Limpa diretório temporário
    import shutil
    try:
        shutil.rmtree(tempdir, ignore_errors=True)
    except Exception:
        pass

    return livro


# ---- EML (Email) ----

def _extrair_eml(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo de email (.eml)"""
    import email
    from email import policy

    conteudo = caminho.read_bytes()
    msg = email.message_from_bytes(conteudo, policy=policy.default)

    # Extrai assunto como título
    assunto = msg.get('Subject', 'Email sem assunto')
    remetente = msg.get('From', 'Desconhecido')
    data = msg.get('Date', '')

    # Extrai corpo do email
    corpo = ""
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == 'text/plain':
                corpo = part.get_content()
                break
            elif content_type == 'text/html' and not corpo:
                html = part.get_content()
                try:
                    from bs4 import BeautifulSoup
                    corpo = BeautifulSoup(html, 'html.parser').get_text(separator='\n\n', strip=True)
                except ImportError:
                    # Fallback: remove tags HTML com regex
                    corpo = re.sub(r'<[^>]+>', '', html)
    else:
        corpo = msg.get_content()

    # Monta texto com cabeçalho
    texto = f"De: {remetente}\nData: {data}\nAssunto: {assunto}\n\n{corpo}"

    livro = _criar_livro_simples(caminho, assunto, texto)
    livro.autor = str(remetente)
    return livro


# ---- MSG (Outlook) ----

def _extrair_msg(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo Outlook (.msg)"""
    try:
        import extract_msg
    except ImportError:
        raise ImportError("Instale extract-msg: pip install extract-msg")

    msg = extract_msg.Message(str(caminho))

    assunto = msg.subject or "Email sem assunto"
    remetente = msg.sender or "Desconhecido"
    corpo = msg.body or ""

    texto = f"De: {remetente}\nAssunto: {assunto}\n\n{corpo}"

    livro = _criar_livro_simples(caminho, assunto, texto)
    livro.autor = str(remetente)

    msg.close()
    return livro


# ---- HTML ----

def _extrair_html(caminho: Path) -> LivroProcessado:
    """Extrai texto de arquivo HTML"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("Instale beautifulsoup4: pip install beautifulsoup4")

    # Tenta diferentes encodings
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            html = caminho.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:
        html = caminho.read_text(encoding='utf-8', errors='replace')

    soup = BeautifulSoup(html, 'html.parser')

    # Remove scripts e estilos
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()

    # Tenta extrair título
    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    if soup.title and soup.title.string:
        titulo = soup.title.string.strip()

    texto = soup.get_text(separator='\n\n', strip=True)
    return _criar_livro_simples(caminho, titulo, texto)


# ---- Imagem (OCR) ----

def _extrair_imagem_ocr(caminho: Path) -> LivroProcessado:
    """Extrai texto de imagem usando OCR (Tesseract)"""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ImportError(
            "Instale pytesseract e Pillow: pip install pytesseract Pillow\n"
            "Também é necessário instalar o Tesseract OCR no sistema:\n"
            "  Windows: https://github.com/UB-Mannheim/tesseract/wiki\n"
            "  Linux: sudo apt install tesseract-ocr tesseract-ocr-por"
        )

    logging.info(f"Aplicando OCR na imagem: {caminho.name}")

    imagem = Image.open(str(caminho))

    # Tenta português primeiro, depois inglês
    try:
        texto = pytesseract.image_to_string(imagem, lang='por')
    except Exception:
        texto = pytesseract.image_to_string(imagem, lang='eng')

    if not texto.strip():
        raise ValueError(f"OCR não conseguiu extrair texto da imagem: {caminho.name}")

    titulo = caminho.stem.replace("_", " ").replace("-", " ")
    livro = _criar_livro_simples(caminho, titulo, texto)
    livro.precisou_ocr = True
    return livro
